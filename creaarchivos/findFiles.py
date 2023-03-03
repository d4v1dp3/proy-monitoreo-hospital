"""
Buscar los archivos de forma recursiva
"""

import os, fnmatch
from docx import Document
from docx.shared import Inches

def find_files(directory, filename):
    for root, dirs, files in os.walk(directory):
        for basename in files:
            if fnmatch.fnmatch(basename, filename):
                encontrado = os.path.join(root, basename)
                return encontrado




def encuentraListaArchivos(fuente, path):
    listaArchivosStr = []
    listaArchivos = []
    with open(fuente,'r') as archFuentes:
        listaArchivosStr= archFuentes.readlines()
    
    print(listaArchivosStr)
    for archivoStr in listaArchivosStr:
        listaArchivos.append(archivoStr.replace("\n",""))
   
    print("#"*10)
    print(listaArchivos)
    return listaArchivos
"""
    listaEncontrados = []
    for archivo in listaArchivos:
        encuentra = find_files(path,archivo)
        print(encuentra)
        # os.system("cat {}".format(encuentra))
        listaEncontrados.append(encuentra)

    return listaEncontrados

"""


#Creando archivo doc con los archivos encontrados
def creaDoc(titulo, listaArchivos, nombreDocx):
    documento = Document()
    documento.add_heading(titulo,0)

    for archivo in listaArchivos:
        with open(archivo,'r') as codeFile:
            contenido = codeFile.readlines()
            #documento.add_heading("Archivo : "+archivo, level=1)
            code =  documento.add_paragraph()
            code.add_run("\n")
            for linea in contenido:
                code.add_run(linea)

    documento.save("{}.docx".format(nombreDocx))




if __name__=="__main__":
    archivos = encuentraListaArchivos("ejbFile.txt",".")
    print(archivos)
    titulo1 = """
      Software del manejador de base de datos del sistema de monitorización del estado de salud de pacientes COVID-19 en hospitales  
    """
    titulo2 = """
        Software del médico del sistema de monitorización del estado de salud
        de pacientes COVID-19 en hospitales
        """

    creaDoc(titulo1, archivos,"CODIGO_FUENTE_SOFTWARE_BASE_DATOS")
    archivos = encuentraListaArchivos("webFile.txt",".")
    creaDoc(titulo2, archivos,"CODIGO_FUENTE_SOFTWARE_")


